import streamlit as st
from auth import init_auth, check_auth, check_role, can_edit_student_info, can_edit_veteran_info, check_page_access
from database import Database
from models import Student, Veteran
from utils import show_success, show_error, apply_theme
from translations import get_text
import pandas as pd
from io import BytesIO
from datetime import datetime
import base64
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from auth import is_search_allowed, is_print_allowed, get_role_based_search_types
import plotly.express as px
import plotly.graph_objects as go
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.units import inch
import os
try:
    from PIL import Image
except ImportError:
    Image = None

def export_student_comprehensive_report(db, student_id):
    """Xuất báo cáo tổng kết toàn diện của học sinh theo mẫu chính thức"""
    from datetime import datetime
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO
    
    # Lấy thông tin học sinh
    student_query = """
        SELECT s.*, c.name as class_name, c.academic_year
        FROM students s
        LEFT JOIN classes c ON s.class_id = c.id
        WHERE s.id = ?
    """
    student_data = db.conn.execute(student_query, (student_id,)).fetchone()
    
    if not student_data:
        return None
    
    # Lấy lịch sử y tế
    medical_records = db.conn.execute("""
        SELECT mr.date, mr.diagnosis, mr.treatment, u.full_name as doctor_name, mr.notes
        FROM medical_records mr
        LEFT JOIN users u ON mr.doctor_id = u.id
        WHERE mr.patient_id = ? AND mr.patient_type = 'student'
        ORDER BY mr.date DESC
    """, (student_id,)).fetchall()
    
    # Lấy tài liệu thực tế đã upload cho học sinh (nếu bảng tồn tại)
    try:
        uploaded_documents = db.conn.execute("""
            SELECT df.file_name, df.description, df.upload_date, u.full_name as uploaded_by, 
                   df.file_type, df.file_data, LENGTH(df.file_data) as file_size
            FROM document_files df
            LEFT JOIN users u ON df.uploaded_by = u.id
            WHERE df.student_id = ?
            ORDER BY df.upload_date DESC
        """, (student_id,)).fetchall()
        print(f"DEBUG: Found {len(uploaded_documents)} documents for student {student_id}")
        for doc in uploaded_documents:
            print(f"DEBUG: Document: {doc[0]} ({doc[4]})")
    except Exception as e:
        print(f"DEBUG: Error loading documents: {e}")
        # Nếu bảng document_files chưa tồn tại, trả về danh sách rỗng
        uploaded_documents = []
    
    # Lấy ghi chú của giáo viên
    teacher_notes = db.conn.execute("""
        SELECT sn.content, sn.note_type, sn.is_important, sn.created_at, 
               u.full_name as teacher_name, c.name as class_name
        FROM student_notes sn
        LEFT JOIN users u ON sn.teacher_id = u.id
        LEFT JOIN classes c ON sn.class_id = c.id
        WHERE sn.student_id = ?
        ORDER BY sn.created_at DESC
    """, (student_id,)).fetchall()
    
    # Lấy lịch sử lớp học chi tiết
    class_history = db.conn.execute("""
        SELECT c.name, c.academic_year, 
               sch.start_date, sch.end_date, sch.notes,
               u.full_name as teacher_name
        FROM student_class_history sch
        JOIN classes c ON sch.class_id = c.id
        LEFT JOIN users u ON c.teacher_id = u.id
        WHERE sch.student_id = ?
        UNION
        SELECT c.name, c.academic_year, 
               COALESCE(s.admission_date, 'Không rõ') as start_date,
               'Đang học' as end_date,
               'Lớp hiện tại' as notes,
               u.full_name as teacher_name
        FROM students s
        JOIN classes c ON s.class_id = c.id
        LEFT JOIN users u ON c.teacher_id = u.id
        WHERE s.id = ? AND s.class_id IS NOT NULL
        ORDER BY start_date DESC
    """, (student_id, student_id)).fetchall()
    

    
    # Tạo document Word
    doc = Document()
    
    # Header
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_para.add_run("ỦY BAN MTTQ VIỆT NAM")
    header_run.bold = True
    header_para.add_run("\n(TÊN ĐƠN VỊ BÁO CÁO)")
    
    header_para2 = doc.add_paragraph()
    header_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run2 = header_para2.add_run("CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    header_run2.bold = True
    header_para2.add_run("\nĐộc lập - Tự do - Hạnh phúc")
    
    # Ngày tháng
    today = datetime.now()
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_para.add_run(f"..., ngày {today.day} tháng {today.month} năm {today.year}")
    
    # Tiêu đề chính
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("BÁO CÁO TỔNG KẾT")
    title_run.bold = True
    
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run(f"CÔNG TÁC MẶT TRẬN NĂM {today.year}")
    subtitle_run.bold = True
    
    # Thông tin học sinh
    doc.add_heading("THÔNG TIN HỌC SINH", level=2)
    
    # Tạo bảng thông tin cơ bản
    info_table = doc.add_table(rows=1, cols=2)
    info_table.style = 'Table Grid'
    
    # Thông tin cơ bản - Fix field mapping based on actual database structure
    basic_info = [
        ("ID Học sinh", student_data[0]),
        ("Họ và tên", student_data[1]),
        ("Ngày sinh", student_data[2] or "Chưa cập nhật"),
        ("Giới tính", student_data[8] or "Chưa cập nhật"),  # Fixed: index 8 for gender
        ("Địa chỉ", student_data[3] or "Chưa cập nhật"),
        ("Email", student_data[4] or "Chưa cập nhật"),
        ("Số điện thoại", student_data[9] or "Chưa cập nhật"),  # Fixed: index 9 for phone
        ("Phụ huynh", student_data[11] or "Chưa cập nhật"),  # Fixed: index 11 for parent_name
        ("Ngày nhập học", student_data[5] or "Chưa cập nhật"),
        ("Năm học", student_data[10] or "Chưa cập nhật"),  # Fixed: index 10 for year
        ("Lớp học hiện tại", student_data[-2] or "Chưa phân lớp"),
        ("Năm học lớp", student_data[-1] or "Chưa cập nhật"),
        ("Số quyết định", student_data[12] or "Chưa cập nhật"),
        ("Thông tin nhà chữ T", student_data[13] or "Chưa cập nhật"),
        ("Sức khỏe khi vào làng", student_data[14] or "Chưa cập nhật"),
        ("Đặc điểm sơ bộ khi vào làng", student_data[15] or "Chưa cập nhật")
    ]
    
    for info_item in basic_info:
        row_cells = info_table.add_row().cells
        row_cells[0].text = info_item[0]
        row_cells[1].text = str(info_item[1])
    

    
    # Lịch sử y tế (luôn hiển thị section)
    doc.add_heading("LỊCH SỬ Y TẾ", level=2)
    
    medical_table = doc.add_table(rows=1, cols=5)
    medical_table.style = 'Table Grid'
    
    # Header
    header_cells = medical_table.rows[0].cells
    header_cells[0].text = "Ngày khám"
    header_cells[1].text = "Chẩn đoán"
    header_cells[2].text = "Điều trị"
    header_cells[3].text = "Bác sĩ"
    header_cells[4].text = "Ghi chú"
    
    if medical_records:
        for record in medical_records:
            row_cells = medical_table.add_row().cells
            row_cells[0].text = str(record[0])[:10] if record[0] else ""  # Date only
            row_cells[1].text = str(record[1]) if record[1] else ""
            row_cells[2].text = str(record[2]) if record[2] else ""
            row_cells[3].text = str(record[3]) if record[3] else ""
            row_cells[4].text = str(record[4]) if record[4] else ""
    else:
        # Add empty row to show no medical records
        row_cells = medical_table.add_row().cells
        row_cells[0].text = "Chưa có lịch sử y tế"
        row_cells[1].text = ""
        row_cells[2].text = ""
        row_cells[3].text = ""
        row_cells[4].text = ""
    
    # Ghi chú về tài liệu đính kèm (không nhúng vào báo cáo)
    if uploaded_documents:
        doc.add_heading("TÀI LIỆU ĐÍNH KÈM", level=2)
        doc.add_paragraph(f"Học sinh này có {len(uploaded_documents)} tài liệu đính kèm:")
        
        for i, doc_record in enumerate(uploaded_documents, 1):
            file_name = doc_record[0] or "Không tên"
            description = doc_record[1] or "Không có mô tả"
            upload_date = doc_record[2] or ""
            
            # Format upload date
            if upload_date:
                try:
                    if isinstance(upload_date, str):
                        parsed_date = datetime.strptime(upload_date, '%Y-%m-%d %H:%M:%S.%f')
                        upload_date = parsed_date.strftime('%d/%m/%Y')
                except:
                    pass
            
            # Thêm thông tin file vào danh sách
            doc.add_paragraph(f"{i}. {file_name} - {description} (Ngày tải: {upload_date})")
        
        doc.add_paragraph()
        note = doc.add_paragraph()
        note.add_run("📌 Ghi chú: ").bold = True
        note.add_run("Các tài liệu trên được xuất riêng cùng với báo cáo này. Vui lòng kiểm tra các file đính kèm.")
    else:
        doc.add_heading("TÀI LIỆU ĐÍNH KÈM", level=2)
        doc.add_paragraph("Chưa có tài liệu nào được tải lên cho học sinh này.")
    
    # Lịch sử lớp học (luôn hiển thị)
    doc.add_heading("LỊCH SỬ LỚP HỌC", level=2)
    
    class_table = doc.add_table(rows=1, cols=6)
    class_table.style = 'Table Grid'
    
    # Header
    header_cells = class_table.rows[0].cells
    header_cells[0].text = "Lớp"
    header_cells[1].text = "Năm học"
    header_cells[2].text = "Ngày bắt đầu"
    header_cells[3].text = "Ngày kết thúc"
    header_cells[4].text = "Giáo viên"
    header_cells[5].text = "Ghi chú"
    
    if class_history:
        for class_record in class_history:
            row_cells = class_table.add_row().cells
            row_cells[0].text = str(class_record[0]) if class_record[0] else ""
            row_cells[1].text = str(class_record[1]) if class_record[1] else ""
            row_cells[2].text = str(class_record[2]) if class_record[2] else ""
            row_cells[3].text = str(class_record[3]) if class_record[3] else "Đang học"
            row_cells[4].text = str(class_record[5]) if len(class_record) > 5 and class_record[5] else ""
            row_cells[5].text = str(class_record[4]) if len(class_record) > 4 and class_record[4] else ""
    else:
        # Thêm hàng trống để báo hiệu chưa có lịch sử
        row_cells = class_table.add_row().cells
        row_cells[0].text = "Chưa có thông tin lớp học"
        row_cells[1].text = ""
        row_cells[2].text = ""
        row_cells[3].text = ""
        row_cells[4].text = ""
        row_cells[5].text = ""

    # Ghi chú của giáo viên (luôn hiển thị section)
    doc.add_heading("GHI CHÚ CỦA GIÁO VIÊN", level=2)
    
    notes_table = doc.add_table(rows=1, cols=6)
    notes_table.style = 'Table Grid'
    
    # Header
    header_cells = notes_table.rows[0].cells
    header_cells[0].text = "Ngày ghi chú"
    header_cells[1].text = "Loại ghi chú"
    header_cells[2].text = "Mức độ"
    header_cells[3].text = "Nội dung"
    header_cells[4].text = "Giáo viên"
    header_cells[5].text = "Lớp"
    
    if teacher_notes:
        for note in teacher_notes:
            row_cells = notes_table.add_row().cells
            row_cells[0].text = str(note[3])[:10] if note[3] else ""  # created_at (date part)
            row_cells[1].text = str(note[1]) if note[1] else ""  # note_type
            row_cells[2].text = "Quan trọng" if note[2] else "Bình thường"  # is_important
            row_cells[3].text = str(note[0]) if note[0] else ""  # content
            row_cells[4].text = str(note[4]) if note[4] else ""  # teacher_name
            row_cells[5].text = str(note[5]) if note[5] else ""  # class_name
    else:
        # Add empty row to show no teacher notes
        row_cells = notes_table.add_row().cells
        row_cells[0].text = "Chưa có ghi chú từ giáo viên"
        row_cells[1].text = ""
        row_cells[2].text = ""
        row_cells[3].text = ""
        row_cells[4].text = ""
        row_cells[5].text = ""
    
    # This section is now handled earlier in the code with uploaded_documents
    
    # Kết luận và đánh giá tổng quan (để viết tay)
    doc.add_heading("TÓM TẮT VÀ ĐÁNH GIÁ TỔNG QUAN", level=2)
    
    summary_para = doc.add_paragraph()
    summary_para.add_run("Đánh giá về sự phát triển của học sinh:")
    summary_para.add_run("\n\n" + "_" * 80)
    summary_para.add_run("\n\n" + "_" * 80)
    summary_para.add_run("\n\n" + "_" * 80)
    summary_para.add_run("\n\n" + "_" * 80)
    summary_para.add_run("\n\n")
    
    summary_para.add_run("Đề xuất hướng dẫn tiếp theo:")
    summary_para.add_run("\n\n" + "_" * 80)
    summary_para.add_run("\n\n" + "_" * 80)
    summary_para.add_run("\n\n" + "_" * 80)
    summary_para.add_run("\n\n")
    
    # Kết luận thống kê
    doc.add_heading("THỐNG KÊ TỔNG QUAN", level=2)
    
    stats_para = doc.add_paragraph()
    stats_para.add_run("Tóm tắt hoạt động và số liệu quan trọng:")
    
    conclusion_list = doc.add_paragraph()
    conclusion_list.add_run(f"• Sức khỏe khi vào làng: {student_data[14] or 'Chưa ghi nhận'}")
    conclusion_list.add_run(f"\n• Đặc điểm sơ bộ khi vào làng: {student_data[15] or 'Chưa ghi nhận'}")
    
    # Thêm tóm tắt ghi chú của giáo viên nếu có
    if teacher_notes:
        important_notes = [note for note in teacher_notes if note[2]]  # is_important = True
        if important_notes:
            conclusion_list.add_run(f"\n• Có {len(important_notes)} ghi chú quan trọng từ giáo viên")
        conclusion_list.add_run(f"\n• Tổng số ghi chú theo dõi: {len(teacher_notes)}")
    
    # Thêm tổng kết y tế và tài liệu
    if medical_records:
        conclusion_list.add_run(f"\n• Số lần khám y tế: {len(medical_records)}")
    if uploaded_documents:
        conclusion_list.add_run(f"\n• Tài liệu đính kèm: {len(uploaded_documents)} file (xuất riêng)")
    if class_history:
        conclusion_list.add_run(f"\n• Số lớp đã học: {len(class_history)}")
    
    # Chữ ký
    doc.add_paragraph("\n\n")
    signature_para = doc.add_paragraph()
    signature_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    signature_para.add_run("Người lập báo cáo")
    signature_para.add_run("\n\n\n")
    signature_para.add_run("(Ký tên và đóng dấu)")
    
    # Lưu file Word
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    # Trả về cả báo cáo Word và danh sách tài liệu đính kèm
    return output, uploaded_documents



def display_student_details(student, db):
    """Hiển thị thông tin chi tiết của học sinh"""
    col1, col2 = st.columns([1, 2])
    
    with col1:
        # Hiển thị ảnh học sinh nếu có
        if student.profile_image:
            try:
                # Kiểm tra kiểu dữ liệu và chuyển đổi nếu cần
                if isinstance(student.profile_image, str):
                    st.warning(get_text('common.invalid_image', 'Dữ liệu ảnh không hợp lệ'))
                    st.image("https://cdn.pixabay.com/photo/2015/10/05/22/37/blank-profile-picture-973460_960_720.png", width=150)
                else:
                    img_data = base64.b64encode(student.profile_image).decode()
                    st.markdown(
                        f'<img src="data:image/jpeg;base64,{img_data}" width="150" style="border-radius: 10px;">',
                        unsafe_allow_html=True
                    )
            except Exception as e:
                st.warning(f"{get_text('common.image_display_error', 'Không thể hiển thị ảnh')}: {str(e)}")
                st.image("https://cdn.pixabay.com/photo/2015/10/05/22/37/blank-profile-picture-973460_960_720.png", width=150)
        else:
            st.image("https://cdn.pixabay.com/photo/2015/10/05/22/37/blank-profile-picture-973460_960_720.png", width=150)
    
    with col2:
        col_a, col_b = st.columns(2)
        not_updated = get_text('common.not_updated', 'Chưa cập nhật')
        
        with col_a:
            st.markdown(f"**🆔 ID Học sinh:** {student.id}")
            st.markdown(f"**🎂 {get_text('profile.birth_date', 'Ngày sinh')}:** {student.birth_date or not_updated}")
            st.markdown(f"**⚧ {get_text('profile.gender', 'Giới tính')}:** {student.gender or not_updated}")
            st.markdown(f"**📱 {get_text('profile.phone', 'Điện thoại')}:** {student.phone or not_updated}")
            st.markdown(f"**📍 {get_text('profile.address', 'Địa chỉ')}:** {student.address or not_updated}")
        
        with col_b:
            # Lấy thông tin lớp học nếu có
            class_info = get_text('profile.no_class', 'Chưa phân lớp')
            if student.class_id:
                class_data = db.get_class(student.class_id)
                if class_data:
                    class_info = f"{class_data.name} ({class_data.academic_year})"
                    
            st.markdown(f"**👨‍🏫 {get_text('profile.class', 'Lớp')}:** {class_info}")
            st.markdown(f"**📧 {get_text('profile.email', 'Email')}:** {student.email or not_updated}")
            st.markdown(f"**👪 {get_text('profile.parent', 'Phụ huynh')}:** {student.parent_name or not_updated}")
            st.markdown(f"**📅 {get_text('profile.admission_date', 'Ngày nhập học')}:** {student.admission_date or not_updated}")
    
    # Hiển thị thông tin hành chính
    st.markdown(f"##### 🏛️ Thông tin hành chính")
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown(f"**📋 Số quyết định:** {getattr(student, 'decision_number', None) or not_updated}")
    
    with col2:
        st.markdown(f"**🏘️ Nhà:** {getattr(student, 'nha_chu_t_info', None) or not_updated}")
    
    # Hiển thị thông tin y tế chi tiết
    st.markdown(f"##### 🏥 Thông tin y tế chi tiết")
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown(f"**💊 Sức khỏe khi vào làng:** {getattr(student, 'health_on_admission', None) or not_updated}")
        st.info("💡 **Lưu ý:** Thông tin chi tiết về hồ sơ y tế được quản lý trong **trang Y tế**")
    
    with col2:
        st.markdown(f"**🔍 Đặc điểm sơ bộ:** {getattr(student, 'initial_characteristics', None) or not_updated}")
    


def display_veteran_details(veteran, db):
    """Hiển thị thông tin chi tiết của cựu chiến binh"""
    col1, col2 = st.columns([1, 2])
    
    with col1:
        # Hiển thị ảnh cựu chiến binh nếu có
        if veteran.profile_image:
            try:
                # Kiểm tra kiểu dữ liệu và chuyển đổi nếu cần
                if isinstance(veteran.profile_image, str):
                    st.warning(get_text('common.invalid_image', 'Dữ liệu ảnh không hợp lệ'))
                    st.image("https://cdn.pixabay.com/photo/2015/10/05/22/37/blank-profile-picture-973460_960_720.png", width=150)
                else:
                    img_data = base64.b64encode(veteran.profile_image).decode()
                    st.markdown(
                        f'<img src="data:image/jpeg;base64,{img_data}" width="150" style="border-radius: 10px;">',
                        unsafe_allow_html=True
                    )
            except Exception as e:
                st.warning(f"{get_text('common.image_display_error', 'Không thể hiển thị ảnh')}: {str(e)}")
                st.image("https://cdn.pixabay.com/photo/2015/10/05/22/37/blank-profile-picture-973460_960_720.png", width=150)
        else:
            st.image("https://cdn.pixabay.com/photo/2015/10/05/22/37/blank-profile-picture-973460_960_720.png", width=150)
    
    with col2:
        col_a, col_b = st.columns(2)
        not_updated = get_text('common.not_updated', 'Chưa cập nhật')
        
        with col_a:
            st.markdown(f"**🆔 ID Bệnh nhân:** {veteran.id}")
            st.markdown(f"**🎂 {get_text('profile.birth_date', 'Ngày sinh')}:** {veteran.birth_date or not_updated}")
            st.markdown(f"**🏅 {get_text('profile.service_period', 'Thời gian phục vụ')}:** {veteran.service_period or not_updated}")
            st.markdown(f"**📍 {get_text('profile.address', 'Địa chỉ')}:** {veteran.address or not_updated}")
        
        with col_b:
            st.markdown(f"**📧 {get_text('profile.email', 'Email')}:** {veteran.email or not_updated}")
            st.markdown(f"**📞 {get_text('profile.contact', 'Liên hệ')}:** {veteran.contact_info or not_updated}")
            st.markdown(f"**🏥 Sức khỏe:** {veteran.health_condition or not_updated}")

def handle_student_edit(student, db):
    """Xử lý chỉnh sửa thông tin học sinh"""
    # Thêm nút đóng ở đầu form
    col_header1, col_header2 = st.columns([4, 1])
    with col_header1:
        st.subheader(f"✏️ {get_text('profile.edit_information', 'Chỉnh sửa thông tin')}")
    with col_header2:
        if st.button("❌ Đóng", key=f"close_edit_{student.id}", help="Đóng form chỉnh sửa"):
            del st.session_state.edit_student_id
            st.rerun()
    
    # Hiển thị ảnh hiện tại và tùy chọn tải lên ảnh mới
    uploaded_image = None
    col_img1, col_img2 = st.columns([1, 2])
    
    with col_img1:
        # Hiển thị ảnh học sinh nếu có
        if student.profile_image:
            try:
                # Kiểm tra kiểu dữ liệu và chuyển đổi nếu cần
                if isinstance(student.profile_image, str):
                    st.warning("Dữ liệu ảnh không hợp lệ")
                    st.image("https://cdn.pixabay.com/photo/2015/10/05/22/37/blank-profile-picture-973460_960_720.png", width=150)
                else:
                    img_data = base64.b64encode(student.profile_image).decode()
                    st.markdown(
                        f'<img src="data:image/jpeg;base64,{img_data}" width="150" style="border-radius: 10px;">',
                        unsafe_allow_html=True
                    )
            except Exception as e:
                st.warning(f"Không thể hiển thị ảnh: {str(e)}")
                st.image("https://cdn.pixabay.com/photo/2015/10/05/22/37/blank-profile-picture-973460_960_720.png", width=150)
        else:
            st.image("https://cdn.pixabay.com/photo/2015/10/05/22/37/blank-profile-picture-973460_960_720.png", width=150)
    
    with col_img2:
        # Tải lên ảnh mới
        uploaded_image = st.file_uploader(get_text("profile.upload_new_image", "Tải lên ảnh mới"), type=["jpg", "jpeg", "png"])
        if uploaded_image is not None:
            # Hiển thị ảnh đã tải lên
            st.image(uploaded_image, width=200)
            
            # Lưu ảnh vào cơ sở dữ liệu khi nhấn nút
            if st.button(get_text("profile.save_image", "Lưu ảnh"), key="save_student_image"):
                try:
                    # Đọc dữ liệu ảnh
                    image_bytes = uploaded_image.getvalue()
                    # Lưu vào cơ sở dữ liệu
                    if db.save_student_image(student.id, image_bytes):
                        show_success("✅ Đã cập nhật ảnh thành công!")
                        st.rerun()  # Reload để hiển thị ảnh mới
                    else:
                        show_error("❌ Không thể cập nhật ảnh")
                except Exception as e:
                    show_error(f"❌ Lỗi khi cập nhật ảnh: {str(e)}")
    
    # Tạo form chỉnh sửa
    with st.form("edit_student_form"):
        col1, col2 = st.columns(2)
        
        with col1:
            full_name = st.text_input(get_text("profile.full_name", "Họ và tên"), value=student.full_name)
            birth_date = st.text_input(get_text("profile.birth_date", "Ngày sinh"), value=student.birth_date or "", 
                                      help=get_text("profile.date_format_help", "Định dạng: YYYY-MM-DD"))
            gender = st.selectbox(get_text("profile.gender", "Giới tính"), 
                                options=[get_text("gender.male", "Nam"), get_text("gender.female", "Nữ"), get_text("gender.other", "Khác")],
                                index=[get_text("gender.male", "Nam"), get_text("gender.female", "Nữ"), get_text("gender.other", "Khác")].index(student.gender) if student.gender in [get_text("gender.male", "Nam"), get_text("gender.female", "Nữ"), get_text("gender.other", "Khác")] else 0)
            phone = st.text_input("Số điện thoại", value=student.phone or "")
            address = st.text_input("Địa chỉ", value=student.address or "")
        
        with col2:
            email = st.text_input("Email", value=student.email or "")
            parent_name = st.text_input("Phụ huynh", value=student.parent_name or "")
            year = st.text_input("Năm học", value=student.year or "")
            
            # Lấy danh sách lớp học
            classes = db.get_classes()
            class_options = [(0, "Không phân lớp")] + [(c.id, f"{c.name} ({c.academic_year})") for c in classes]
            
            # Tìm vị trí của lớp học hiện tại trong danh sách
            selected_index = 0
            for i, (class_id, _) in enumerate(class_options):
                if student.class_id == class_id:
                    selected_index = i
                    break
            
            class_id = st.selectbox(
                "Lớp", 
                options=range(len(class_options)),
                format_func=lambda i: class_options[i][1],
                index=selected_index
            )
            class_id = class_options[class_id][0] if class_id is not None else None
            
            # Convert class_id = 0 to None
            if class_id == 0:
                class_id = None
            
            # Đổi định dạng ngày nhập học từ text sang date_input
            admission_date = None
            try:
                if student.admission_date:
                    admission_date = datetime.strptime(student.admission_date, "%Y-%m-%d").date()
            except:
                pass
            
            admission_date = st.date_input("Ngày nhập học", value=admission_date or datetime.now().date())
        
        # Thông tin hành chính
        st.subheader("🏛️ Thông tin hành chính")
        col1, col2 = st.columns(2)
        
        with col1:
            decision_number = st.text_input("Số quyết định", value=getattr(student, 'decision_number', '') or "")
        
        with col2:
            nha_chu_t_info = st.text_input("Nhà", value=getattr(student, 'nha_chu_t_info', '') or "")
        
        # Thông tin y tế chi tiết
        st.subheader("🏥 Thông tin y tế chi tiết")
        col1, col2 = st.columns(2)
        
        with col1:
            health_on_admission = st.text_area("Sức khỏe khi vào làng", value=getattr(student, 'health_on_admission', '') or "")
        
        with col2:
            initial_characteristics = st.text_area("Đặc điểm sơ bộ khi vào làng", value=getattr(student, 'initial_characteristics', '') or "", 
                                                  help="Mô tả đặc điểm và tình trạng ban đầu của bệnh nhân khi vào làng")
        

            
        if st.form_submit_button("💾 Lưu thay đổi"):
            student_data = {
                "full_name": full_name,
                "birth_date": birth_date,
                "gender": gender,
                "phone": phone,
                "address": address,
                "email": email,
                "parent_name": parent_name,
                "year": year,
                "admission_date": admission_date.strftime("%Y-%m-%d"),
                "class_id": class_id,
                "decision_number": decision_number,
                "nha_chu_t_info": nha_chu_t_info,
                "health_on_admission": health_on_admission,
                "initial_characteristics": initial_characteristics
            }
            
            try:
                if db.update_student(student.id, student_data):
                    # Ghi nhận lịch sử thay đổi lớp học nếu lớp học thay đổi
                    if student.class_id != class_id:
                        db.update_student_class(student.id, class_id)
                    
                    show_success("✅ Đã cập nhật thông tin học sinh thành công!")
                    st.rerun()
                else:
                    show_error("❌ Không thể cập nhật thông tin học sinh")
            except Exception as e:
                show_error(f"❌ Lỗi khi cập nhật: {str(e)}")

def handle_veteran_edit(veteran, db):
    """Xử lý chỉnh sửa thông tin cựu chiến binh"""
    st.subheader("✏️ Chỉnh sửa thông tin")
    
    # Hiển thị ảnh hiện tại và tùy chọn tải lên ảnh mới
    uploaded_image = None
    col_img1, col_img2 = st.columns([1, 2])
    
    with col_img1:
        # Hiển thị ảnh cựu chiến binh nếu có
        if veteran.profile_image:
            try:
                # Kiểm tra kiểu dữ liệu và chuyển đổi nếu cần
                if isinstance(veteran.profile_image, str):
                    st.warning("Dữ liệu ảnh không hợp lệ")
                    st.image("https://cdn.pixabay.com/photo/2015/10/05/22/37/blank-profile-picture-973460_960_720.png", width=150)
                else:
                    img_data = base64.b64encode(veteran.profile_image).decode()
                    st.markdown(
                        f'<img src="data:image/jpeg;base64,{img_data}" width="150" style="border-radius: 10px;">',
                        unsafe_allow_html=True
                    )
            except Exception as e:
                st.warning(f"Không thể hiển thị ảnh: {str(e)}")
                st.image("https://cdn.pixabay.com/photo/2015/10/05/22/37/blank-profile-picture-973460_960_720.png", width=150)
        else:
            st.image("https://cdn.pixabay.com/photo/2015/10/05/22/37/blank-profile-picture-973460_960_720.png", width=150)
    
    with col_img2:
        # Tải lên ảnh mới
        uploaded_image = st.file_uploader(get_text("profile.upload_new_image", "Tải lên ảnh mới"), type=["jpg", "jpeg", "png"], key="upload_veteran_image")
        if uploaded_image is not None:
            # Hiển thị ảnh đã tải lên
            st.image(uploaded_image, width=200)
            
            # Lưu ảnh vào cơ sở dữ liệu khi nhấn nút
            if st.button(get_text("profile.save_image", "Lưu ảnh"), key="save_veteran_image"):
                try:
                    # Đọc dữ liệu ảnh
                    image_bytes = uploaded_image.getvalue()
                    # Lưu vào cơ sở dữ liệu
                    if db.save_veteran_image(veteran.id, image_bytes):
                        show_success("✅ Đã cập nhật ảnh thành công!")
                        st.rerun()  # Reload để hiển thị ảnh mới
                    else:
                        show_error("❌ Không thể cập nhật ảnh")
                except Exception as e:
                    show_error(f"❌ Lỗi khi cập nhật ảnh: {str(e)}")
    
    # Tạo form chỉnh sửa
    with st.form("edit_veteran_form"):
        col1, col2 = st.columns(2)
        
        with col1:
            full_name = st.text_input("Họ và tên", value=veteran.full_name)
            birth_date = st.text_input("Ngày sinh", value=veteran.birth_date or "", 
                                      help="Định dạng: YYYY-MM-DD")
            service_period = st.text_input("Thời gian phục vụ", value=veteran.service_period or "")
            address = st.text_input("Địa chỉ", value=veteran.address or "")
        
        with col2:
            health_condition = st.text_input("Sức khỏe", value=veteran.health_condition or "")
            email = st.text_input("Email", value=veteran.email or "")
            contact_info = st.text_input("Thông tin liên hệ", value=veteran.contact_info or "")
            
        # Đặc điểm sơ bộ
        st.subheader("📋 Đặc điểm sơ bộ")
        initial_characteristics = st.text_area("Đặc điểm sơ bộ khi vào làng", 
                                              value=getattr(veteran, 'initial_characteristics', '') or "",
                                              help="Mô tả đặc điểm và tình trạng ban đầu của cựu chiến binh khi vào làng")
            
        if st.form_submit_button("💾 Lưu thay đổi"):
            veteran_data = {
                "full_name": full_name,
                "birth_date": birth_date,
                "service_period": service_period,
                "address": address,
                "health_condition": health_condition,
                "email": email,
                "contact_info": contact_info,
                "initial_characteristics": initial_characteristics
            }
            
            try:
                if db.update_veteran(veteran.id, veteran_data):
                    show_success("✅ Đã cập nhật thông tin cựu chiến binh thành công!")
                    st.rerun()
                else:
                    show_error("❌ Không thể cập nhật thông tin cựu chiến binh")
            except Exception as e:
                show_error(f"❌ Lỗi khi cập nhật: {str(e)}")

def add_new_student(db):
    """Thêm học sinh mới"""
    st.subheader("➕ Thêm học sinh mới")
    
    # Tải lên ảnh cho học sinh mới
    uploaded_image = st.file_uploader("Tải lên ảnh học sinh", type=["jpg", "jpeg", "png"], key="upload_new_student_image")
    if uploaded_image is not None:
        st.image(uploaded_image, width=200)
        st.info("Ảnh sẽ được lưu sau khi bạn thêm học sinh")
    
    with st.form("add_student_form"):
        # Thông tin cơ bản
        col1, col2 = st.columns(2)
        with col1:
            full_name = st.text_input("Họ và tên", key="add_full_name")
            birth_date = st.text_input("Ngày sinh", key="add_birth_place", 
                                     help="Định dạng: YYYY-MM-DD")
            gender = st.selectbox("Giới tính", ["Nam", "Nữ", "Khác"], key="add_gender")
            phone = st.text_input("Số điện thoại", key="add_phone")
            address = st.text_input("Địa chỉ thường trú", key="add_address")
        
        with col2:
            email = st.text_input("Email", key="add_email")
            parent_name = st.text_input("Phụ huynh", key="add_parent_name")
            year = st.text_input("Năm học", key="add_year")
            admission_date = st.date_input("Ngày nhập học", key="add_admission_date")

            # Lấy danh sách lớp học
            classes = db.get_classes()
            class_options = [(0, "Không phân lớp")] + [(c.id, f"{c.name} ({c.academic_year})") for c in classes]
            
            class_id = st.selectbox(
                "Lớp",
                options=range(len(class_options)),
                format_func=lambda i: class_options[i][1],
                key="add_class_id"
            )
            class_id = class_options[class_id][0] if class_id is not None else None
            
            # Convert class_id = 0 to None
            if class_id == 0:
                class_id = None

        # Thông tin hành chính
        st.subheader("🏛️ Thông tin hành chính")
        col1, col2 = st.columns(2)
        
        with col1:
            decision_number = st.text_input(
                "Số quyết định",
                key="add_decision_number",
                help="Số quyết định nhập học"
            )
        
        with col2:
            nha_chu_t_info = st.text_area(
                "Nhà",
                key="add_nha_chu_t_info",
                help="Thông tin liên quan đến nhà chữ T"
            )

        # Tình trạng y tế mở rộng
        st.subheader("🏥 Thông tin y tế chi tiết")
        col1, col2 = st.columns(2)
        
        with col1:
            health_on_admission = st.text_area(
                "Tình trạng sức khỏe khi vào làng",
                key="add_health_on_admission",
                help="Mô tả tình trạng sức khỏe ban đầu khi nhập học"
            )
        
        with col2:
            initial_characteristics = st.text_area(
                "Đặc điểm sơ bộ khi vào làng",
                key="add_initial_characteristics",
                help="Mô tả đặc điểm và tình trạng ban đầu của bệnh nhân khi vào làng"
            )



        if st.form_submit_button("Thêm học sinh"):
            student_data = {
                "full_name": full_name,
                "birth_date": birth_date,
                "gender": gender,
                "phone": phone,
                "address": address,
                "email": email,
                "parent_name": parent_name,
                "year": year,
                "admission_date": admission_date.strftime("%Y-%m-%d"),

                "class_id": class_id,
                "decision_number": decision_number,
                "nha_chu_t_info": nha_chu_t_info,
                "health_on_admission": health_on_admission,
                "initial_characteristics": initial_characteristics
            }
            
            if not full_name:
                show_error("Vui lòng nhập họ và tên học sinh")
            else:
                try:
                    student_id = db.add_student(student_data)
                    if student_id:
                        # Nếu có lớp học, cập nhật lịch sử lớp học
                        if class_id:
                            db.update_student_class(student_id, class_id)
                        
                        # Nếu có tải lên ảnh, lưu ảnh vào cơ sở dữ liệu
                        if 'upload_new_student_image' in st.session_state and st.session_state.upload_new_student_image is not None:
                            uploaded_image = st.session_state.upload_new_student_image
                            image_bytes = uploaded_image.getvalue()
                            if db.save_student_image(student_id, image_bytes):
                                show_success(f"Đã thêm học sinh {full_name} và lưu ảnh thành công!")
                            else:
                                show_success(f"Đã thêm học sinh {full_name} thành công, nhưng không thể lưu ảnh!")
                        else:
                            show_success(f"Đã thêm học sinh {full_name} thành công!")
                            
                        st.rerun()
                    else:
                        show_error("Không thể thêm học sinh")
                except Exception as e:
                    show_error(f"Lỗi khi thêm học sinh: {str(e)}")

def add_new_veteran(db):
    """Thêm cựu chiến binh mới"""
    st.subheader("➕ Thêm cựu chiến binh mới")
    
    # Tải lên ảnh cho cựu chiến binh mới
    uploaded_image = st.file_uploader("Tải lên ảnh cựu chiến binh", type=["jpg", "jpeg", "png"], key="upload_new_veteran_image")
    if uploaded_image is not None:
        st.image(uploaded_image, width=200)
        st.info("Ảnh sẽ được lưu sau khi bạn thêm cựu chiến binh")
    
    with st.form("add_veteran_form"):
        col1, col2 = st.columns(2)
        
        with col1:
            full_name = st.text_input("Họ và tên", key="add_vet_name")
            birth_date = st.text_input("Ngày sinh", key="add_vet_birth", 
                                     help="Định dạng: YYYY-MM-DD")
            service_period = st.text_input("Thời gian phục vụ", key="add_vet_service")
            address = st.text_input("Địa chỉ", key="add_vet_address")
        
        with col2:
            health_condition = st.text_input("Sức khỏe", key="add_vet_health")
            email = st.text_input("Email", key="add_vet_email")
            contact_info = st.text_input("Thông tin liên hệ", key="add_vet_contact")
            
        # Đặc điểm sơ bộ
        st.subheader("📋 Đặc điểm sơ bộ")
        initial_characteristics = st.text_area("Đặc điểm sơ bộ khi vào làng", 
                                              key="add_vet_initial_characteristics",
                                              help="Mô tả đặc điểm và tình trạng ban đầu của cựu chiến binh khi vào làng")
            
        if st.form_submit_button("Thêm cựu chiến binh"):
            if not full_name:
                show_error("Vui lòng nhập họ và tên")
            else:
                veteran_data = {
                    "full_name": full_name,
                    "birth_date": birth_date,
                    "service_period": service_period,
                    "address": address,
                    "health_condition": health_condition,
                    "email": email,
                    "contact_info": contact_info,
                    "initial_characteristics": initial_characteristics
                }
                
                try:
                    veteran_id = db.add_veteran(veteran_data)
                    if veteran_id:
                        # Nếu có tải lên ảnh, lưu ảnh vào cơ sở dữ liệu
                        if 'upload_new_veteran_image' in st.session_state and st.session_state.upload_new_veteran_image is not None:
                            uploaded_image = st.session_state.upload_new_veteran_image
                            image_bytes = uploaded_image.getvalue()
                            if db.save_veteran_image(veteran_id, image_bytes):
                                show_success(f"Đã thêm cựu chiến binh {full_name} và lưu ảnh thành công!")
                            else:
                                show_success(f"Đã thêm cựu chiến binh {full_name} thành công, nhưng không thể lưu ảnh!")
                        else:
                            show_success(f"Đã thêm cựu chiến binh {full_name} thành công!")
                            
                        st.rerun()
                    else:
                        show_error("Không thể thêm cựu chiến binh")
                except Exception as e:
                    show_error(f"Lỗi khi thêm cựu chiến binh: {str(e)}")

def render():
    # Initialize authentication first
    init_auth()
    
    # Apply theme from session state
    # apply_theme()  # Disabled - using Streamlit defaults
    
    # Set current page for role checking
    st.session_state.current_page = "02_Quản_lý_hồ_sơ"
    
    check_auth()
    
    # Check if user has access to this page dynamically
    if not check_page_access('02_Quản_lý_hồ_sơ'):
        st.error("Bạn không có quyền truy cập trang này.")
        st.stop()
        return
    
    db = Database()
    
    st.title(f"👥 {get_text('profile.title', 'Quản Lý Hồ Sơ')}")
    
    # Tạo tabs chính cho quản lý hồ sơ, tài liệu, thống kê và xuất dữ liệu
    main_tabs = st.tabs([
        "👥 Quản lý hồ sơ",
        "📁 Quản lý tài liệu",
        "📊 Thống kê và báo cáo", 
        "📤 Xuất dữ liệu"
    ])
    
    with main_tabs[0]:
        # Tạo tabs cho học sinh và cựu chiến binh
        entity_type = st.radio(
            get_text('profile.select_type', 'Chọn loại hồ sơ'), 
            [get_text('common.student', 'Học sinh'), get_text('common.veteran', 'Cựu chiến binh')], 
            horizontal=True
        )
        
        # Thêm tìm kiếm nâng cao với dropdown
        with st.expander("🔍 Tìm kiếm nâng cao", expanded=False):
            if entity_type == get_text('common.student', 'Học sinh'):
                render_student_advanced_search(db)
            else:
                render_veteran_advanced_search(db)
        
        if entity_type == get_text('common.student', 'Học sinh'):
            student_tabs = st.tabs([
                f"🔍 {get_text('common.list', 'Danh sách')}", 
                f"➕ {get_text('common.add_new', 'Thêm mới')}"
            ])
            
            with student_tabs[0]:
                st.subheader(f"📋 {get_text('profile.student_list', 'Danh sách học sinh')}")
                
                # Thêm ô tìm kiếm
                search_query = st.text_input(
                    f"🔍 {get_text('common.search_by_name', 'Tìm kiếm theo tên')}", 
                    key="search_student"
                )
                
                # Lấy danh sách học sinh
                students = db.get_students()
                
                # Lọc học sinh theo tìm kiếm nếu có
                if search_query:
                    search_query = search_query.lower()
                    students = [s for s in students if search_query in s.full_name.lower()]
                
                # Hiển thị số lượng học sinh
                st.info(f"📊 {get_text('common.total', 'Tổng số')}: {len(students)} {get_text('common.student', 'học sinh')}")
                
                # Hiển thị danh sách học sinh trong các expanders
                for student in students:
                    # Kiểm tra nếu đang chỉnh sửa học sinh này
                    is_editing = 'edit_student_id' in st.session_state and st.session_state.edit_student_id == student.id
                    
                    with st.expander(f"👨‍🎓 {student.full_name}" + (" ✏️ (Đang chỉnh sửa)" if is_editing else ""), expanded=is_editing):
                        # Thông tin chi tiết học sinh
                        display_student_details(student, db)
                        
                        # Nút chỉnh sửa và xuất báo cáo
                        col1, col2 = st.columns(2)
                        with col1:
                            if can_edit_student_info():
                                if not is_editing:
                                    if st.button(f"✏️ {get_text('common.edit', 'Chỉnh sửa')}", key=f"edit_student_{student.id}"):
                                        st.session_state.edit_student_id = student.id
                                        st.rerun()
                                else:
                                    if st.button(f"❌ {get_text('common.cancel_edit', 'Hủy chỉnh sửa')}", key=f"cancel_edit_{student.id}"):
                                        del st.session_state.edit_student_id
                                        st.rerun()
                        
                        with col2:
                            if st.button("📄 Xuất báo cáo Word", key=f"export_report_{student.id}"):
                                try:
                                    word_file, documents = export_student_comprehensive_report(db, student.id)
                                    if word_file:
                                        st.success("Báo cáo đã được tạo thành công!")
                                        
                                        # Tải xuống báo cáo Word
                                        st.download_button(
                                            label="📥 Tải xuống báo cáo Word",
                                            data=word_file,
                                            file_name=f"bao_cao_tong_ket_{student.full_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx",
                                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                            key=f"download_report_{student.id}"
                                        )
                                        
                                        # Xuất tài liệu đính kèm nếu có
                                        if documents:
                                            st.info(f"📎 Tìm thấy {len(documents)} tài liệu đính kèm. Tải xuống từng file:")
                                            
                                            # Tạo cột để hiển thị các nút tải xuống tài liệu
                                            cols = st.columns(min(len(documents), 3))
                                            for i, doc in enumerate(documents):
                                                col_idx = i % 3
                                                with cols[col_idx]:
                                                    file_name = doc[0] or "unknown_file"
                                                    file_data = doc[5]  # file_data BLOB
                                                    file_type = doc[4] or ""
                                                    
                                                    # Tạo mime type từ file type
                                                    mime_type = "application/octet-stream"
                                                    if file_type:
                                                        if "pdf" in file_type.lower():
                                                            mime_type = "application/pdf"
                                                        elif "word" in file_type.lower() or "docx" in file_type.lower():
                                                            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                                        elif "excel" in file_type.lower() or "xlsx" in file_type.lower():
                                                            mime_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                                                        elif "image" in file_type.lower() or any(ext in file_type.lower() for ext in ['jpg', 'jpeg', 'png', 'gif']):
                                                            mime_type = file_type
                                                    
                                                    st.download_button(
                                                        label=f"📎 {file_name}",
                                                        data=file_data,
                                                        file_name=file_name,
                                                        mime=mime_type,
                                                        key=f"download_doc_{student.id}_{i}",
                                                        help=f"Tải xuống: {doc[1] or 'Không có mô tả'}"
                                                    )
                                        else:
                                            st.info("📝 Báo cáo không có tài liệu đính kèm")
                                    else:
                                        show_error("Không thể tạo báo cáo")
                                except Exception as e:
                                    show_error(f"Lỗi khi tạo báo cáo: {str(e)}")
                        
                        # Hiển thị form chỉnh sửa ngay trong expander nếu đang chỉnh sửa
                        if is_editing:
                            st.divider()
                            handle_student_edit(student, db)
        
            with student_tabs[1]:
                # Kiểm tra quyền thêm học sinh mới
                if can_edit_student_info():
                    add_new_student(db)
                else:
                    st.warning(get_text("profile.student_add_permission_denied", "Bạn không có quyền thêm học sinh mới"))
    
        else:  # Cựu chiến binh
            veteran_tabs = st.tabs([
                f"🔍 {get_text('common.list', 'Danh sách')}", 
                f"➕ {get_text('common.add_new', 'Thêm mới')}"
            ])
            
            with veteran_tabs[0]:
                st.subheader(f"📋 {get_text('profile.veteran_list', 'Danh sách cựu chiến binh')}")
            
            # Thêm ô tìm kiếm
            search_query = st.text_input(
                f"🔍 {get_text('common.search_by_name', 'Tìm kiếm theo tên')}", 
                key="search_veteran"
            )
            
            # Lấy danh sách cựu chiến binh
            veterans = db.get_veterans()
            
            # Lọc cựu chiến binh theo tìm kiếm nếu có
            if search_query:
                search_query = search_query.lower()
                veterans = [v for v in veterans if search_query in v.full_name.lower()]
            
            # Hiển thị số lượng cựu chiến binh
            st.info(f"📊 {get_text('common.total', 'Tổng số')}: {len(veterans)} {get_text('common.veteran', 'cựu chiến binh')}")
            
            # Hiển thị danh sách cựu chiến binh trong các expanders
            for veteran in veterans:
                with st.expander(f"🎖️ {veteran.full_name}"):
                    # Thông tin chi tiết cựu chiến binh
                    display_veteran_details(veteran, db)
                    
                    # Nút chỉnh sửa
                    if can_edit_veteran_info():
                        if st.button(f"✏️ {get_text('common.edit', 'Chỉnh sửa')}", key=f"edit_veteran_{veteran.id}"):
                            st.session_state.edit_veteran_id = veteran.id
                            st.rerun()
            
            # Hiển thị form chỉnh sửa nếu đã chọn cựu chiến binh
            if hasattr(st.session_state, 'edit_veteran_id'):
                for veteran in veterans:
                    if veteran.id == st.session_state.edit_veteran_id:
                        handle_veteran_edit(veteran, db)
                        break
                
                # Nút hủy chỉnh sửa
                if st.button(f"❌ {get_text('common.cancel_edit', 'Hủy chỉnh sửa')}", key="cancel_veteran_edit"):
                    del st.session_state.edit_veteran_id
                    st.rerun()
        
            with veteran_tabs[1]:
                # Kiểm tra quyền thêm cựu chiến binh mới
                if can_edit_veteran_info():
                    add_new_veteran(db)
                else:
                    st.warning(get_text("profile.veteran_add_permission_denied", "Bạn không có quyền thêm cựu chiến binh mới"))
    
    with main_tabs[1]:
        # Quản lý tài liệu
        render_document_management_section(db)
    
    with main_tabs[2]:
        # Thống kê và báo cáo
        render_statistics_section(db)
    
    with main_tabs[3]:
        # Xuất dữ liệu
        render_export_section(db)

def render_statistics_section(db):
    """Render statistics and analytics section"""
    st.subheader("📊 Thống kê và phân tích dữ liệu")
    
    # Import required modules
    import plotly.express as px
    import plotly.graph_objects as go
    from datetime import datetime, timedelta
    
    # Overview statistics
    col1, col2, col3, col4 = st.columns(4)
    
    total_students = len(db.get_students())
    total_veterans = len(db.get_veterans())
    
    medical_records = db.conn.execute("SELECT COUNT(*) FROM medical_records").fetchone()[0]
    psych_evals = db.conn.execute("SELECT COUNT(*) FROM psychological_evaluations").fetchone()[0]
    
    with col1:
        st.metric("Tổng số học sinh", total_students)
    with col2:
        st.metric("Tổng số cựu chiến binh", total_veterans)
    with col3:
        st.metric("Số hồ sơ y tế", medical_records)
    with col4:
        st.metric("Số đánh giá tâm lý", psych_evals)
    
    # Charts section
    st.subheader("📈 Biểu đồ thống kê")
    
    chart_tabs = st.tabs([
        "📊 Phân bố lớp học"
    ])
    
    with chart_tabs[0]:
        # Class distribution chart
        data = db.conn.execute("""
            SELECT c.name, COUNT(s.id) as student_count
            FROM classes c
            LEFT JOIN students s ON c.id = s.class_id
            GROUP BY c.id, c.name
            ORDER BY student_count DESC
        """).fetchall()
        
        if data:
            df = pd.DataFrame(data, columns=['class_name', 'count'])
            fig = px.bar(df, x='class_name', y='count',
                        title='Phân Bố Học Sinh Theo Lớp')
            fig.update_layout(xaxis_tickangle=45)
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.warning("Không có dữ liệu lớp học")





def export_student_list_integrated(students, format_type='excel'):
    """Export student list from search results"""
    if not students:
        st.warning("Không có học sinh nào để xuất")
        return
    
    db = Database()
    
    # Prepare data for export
    data = []
    for student in students:
        class_name = "Chưa phân lớp"
        if student.class_id:
            class_info = db.get_class(student.class_id)
            if class_info:
                class_name = class_info.name
        
        student_data = {
            "ID": student.id,
            "Họ và tên": student.full_name,
            "Ngày sinh": student.birth_date or "Chưa cập nhật",
            "Giới tính": student.gender or "Chưa cập nhật",
            "Địa chỉ": student.address or "Chưa cập nhật",
            "Email": student.email or "Chưa cập nhật",
            "Số điện thoại": student.phone or "Chưa cập nhật",
            "Lớp": class_name,
            "Ngày nhập học": student.admission_date or "Chưa cập nhật",

        }
        data.append(student_data)
    
    df = pd.DataFrame(data)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    if format_type == 'excel':
        output = BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name="Danh sách học sinh")
        
        excel_data = output.getvalue()
        filename = f"Danh_sach_hoc_sinh_{timestamp}.xlsx"
        
        st.download_button(
            label="📥 Tải xuống danh sách (Excel)",
            data=excel_data,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key="download_student_list_integrated_excel"
        )
        st.success("Đã tạo file Excel thành công!")
        
    elif format_type == 'csv':
        csv = df.to_csv(index=False).encode('utf-8')
        filename = f"Danh_sach_hoc_sinh_{timestamp}.csv"
        
        st.download_button(
            label="📥 Tải xuống danh sách (CSV)",
            data=csv,
            file_name=filename,
            mime="text/csv",
            key="download_student_list_integrated_csv"
        )
        st.success("Đã tạo file CSV thành công!")

def render_student_advanced_search(db):
    """Render advanced search for students"""
    # Check search permissions
    if not is_search_allowed():
        st.error("Bạn không có quyền sử dụng chức năng tìm kiếm")
        return
    
    with st.form("student_advanced_search"):
        st.subheader("🔍 Tìm kiếm học sinh nâng cao")
        
        col1, col2 = st.columns(2)
        with col1:
            name_query = st.text_input("Tên học sinh", key="adv_student_name")
            address = st.text_input("Địa chỉ", key="adv_student_address")
            email = st.text_input("Email", key="adv_student_email")
            phone = st.text_input("Số điện thoại", key="adv_student_phone")
            
        with col2:
            gender = st.selectbox("Giới tính", ["Tất cả", "Nam", "Nữ"], key="adv_student_gender")
            year = st.text_input("Năm", key="adv_student_year")
            parent_name = st.text_input("Phụ huynh", key="adv_student_parent")
        
        # Class selection
        class_options = [(None, "Tất cả lớp")] + [(c.id, c.name) for c in db.get_classes()]
        selected_class = st.selectbox(
            "Lớp",
            options=range(len(class_options)),
            format_func=lambda i: class_options[i][1],
            key="adv_student_class"
        )
        class_id = class_options[selected_class][0] if selected_class > 0 else None
        

        
        search_button = st.form_submit_button("🔍 Tìm kiếm", type="primary")
    
    # Handle search results outside the form
    if search_button:
        students = search_students_advanced(db, name_query, address, email, phone, gender, year, parent_name, 
                                         class_id)
        st.session_state.adv_search_results_students = students
    
    # Display results outside the form
    if 'adv_search_results_students' in st.session_state and st.session_state.adv_search_results_students:
        students = st.session_state.adv_search_results_students
        st.success(f"🔎 Tìm thấy {len(students)} học sinh")
        
        # Export options
        col1, col2 = st.columns(2)
        with col1:
            if st.button("📊 Xuất Excel", key="export_students_excel_adv"):
                export_student_list_integrated(students, 'excel')
        with col2:
            if st.button("📄 Xuất CSV", key="export_students_csv_adv"):
                export_student_list_integrated(students, 'csv')
        
        # Display results
        for i, student in enumerate(students):
            st.divider()
            st.subheader(f"👨‍🎓 ID {student.id} - {student.full_name}")
            
            col1, col2 = st.columns(2)
            with col1:
                st.markdown(f"**🆔 ID Học sinh:** {student.id}")
                st.markdown(f"**🎂 Ngày sinh:** {student.birth_date or 'Chưa cập nhật'}")
                st.markdown(f"**📍 Địa chỉ:** {student.address or 'Chưa cập nhật'}")
                st.markdown(f"**📧 Email:** {student.email or 'Chưa cập nhật'}")
                st.markdown(f"**📞 Điện thoại:** {student.phone or 'Chưa cập nhật'}")
            with col2:
                st.markdown(f"**🏥 Sức khỏe khi vào làng:** {getattr(student, 'health_on_admission', '') or 'Chưa cập nhật'}")
                st.markdown(f"**🔍 Đặc điểm sơ bộ:** {getattr(student, 'initial_characteristics', '') or 'Chưa cập nhật'}")
                st.markdown(f"**🏠 Nhà:** {getattr(student, 'nha_chu_t_info', '') or 'Chưa cập nhật'}")
            
            if st.button(f"📄 Xuất báo cáo Word", key=f"export_word_adv_{student.id}"):
                word_file = export_student_comprehensive_report(db, student.id)
                if word_file:
                    st.success("Báo cáo đã được tạo thành công!")
                    st.download_button(
                        label="📥 Tải xuống báo cáo Word",
                        data=word_file,
                        file_name=f"bao_cao_tong_ket_{student.full_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"download_word_adv_{student.id}"
                    )
                else:
                    st.error("Không thể tạo báo cáo")
    elif search_button:
        st.info("Không tìm thấy học sinh nào phù hợp")

def render_veteran_advanced_search(db):
    """Render advanced search for veterans"""
    # Check search permissions
    if not is_search_allowed():
        st.error("Bạn không có quyền sử dụng chức năng tìm kiếm")
        return
    
    with st.form("veteran_advanced_search"):
        st.subheader("🔍 Tìm kiếm cựu chiến binh nâng cao")
        
        col1, col2 = st.columns(2)
        with col1:
            name_query = st.text_input("Tên cựu chiến binh", key="adv_veteran_name")
            address = st.text_input("Địa chỉ", key="adv_veteran_address")
            email = st.text_input("Email", key="adv_veteran_email")
        with col2:
            gender = st.selectbox("Giới tính", ["Tất cả", "Nam", "Nữ"], key="adv_veteran_gender")
            health_status = st.selectbox(
                "Đặc điểm ban đầu",
                ["Tất cả", "Tốt", "Bình thường", "Cần chú ý"],
                key="adv_veteran_health"
            )
        
        search_button = st.form_submit_button("🔍 Tìm kiếm", type="primary")
    
    # Handle search results outside the form
    if search_button:
        veterans = search_veterans_advanced(db, name_query, address, email, gender, health_status)
        st.session_state.adv_search_results_veterans = veterans
    
    # Display results outside the form
    if 'adv_search_results_veterans' in st.session_state and st.session_state.adv_search_results_veterans:
        veterans = st.session_state.adv_search_results_veterans
        st.success(f"🔎 Tìm thấy {len(veterans)} cựu chiến binh")
        
        # Display results
        for i, veteran in enumerate(veterans):
            st.divider()
            st.subheader(f"🎖️ ID {veteran.id} - {veteran.full_name}")
            
            col1, col2 = st.columns(2)
            with col1:
                st.markdown(f"**🆔 ID Bệnh nhân:** {veteran.id}")
                st.markdown(f"**🎂 Ngày sinh:** {veteran.birth_date or 'Chưa cập nhật'}")
                st.markdown(f"**📍 Địa chỉ:** {veteran.address or 'Chưa cập nhật'}")
            with col2:
                st.markdown(f"**📧 Email:** {veteran.email or 'Chưa cập nhật'}")
                st.markdown(f"**🏥 Đặc điểm ban đầu:** {veteran.health_condition or 'Chưa cập nhật'}")
                st.markdown(f"**⚧ Giới tính:** {veteran.gender or 'Chưa cập nhật'}")
    elif search_button:
        st.info("Không tìm thấy cựu chiến binh nào phù hợp")

def search_students_advanced(db, name_query, address, email, phone, gender, year, parent_name, 
                           class_id):
    """Advanced search for students"""
    conditions = []
    params = []
    
    if name_query:
        conditions.append("s.full_name LIKE ?")
        params.append(f"%{name_query}%")
    
    if address:
        conditions.append("s.address LIKE ?")
        params.append(f"%{address}%")
        
    if email:
        conditions.append("s.email LIKE ?")
        params.append(f"%{email}%")
        
    if phone:
        conditions.append("s.phone LIKE ?")
        params.append(f"%{phone}%")
        
    if gender and gender != "Tất cả":
        conditions.append("s.gender = ?")
        params.append(gender)
        
    if year:
        conditions.append("s.year LIKE ?")
        params.append(f"%{year}%")
        
    if parent_name:
        conditions.append("s.parent_name LIKE ?")
        params.append(f"%{parent_name}%")
        
    if class_id:
        conditions.append("s.class_id = ?")
        params.append(class_id)
        

    
    where_clause = " AND ".join(conditions) if conditions else "1=1"
    
    query = f"""
        SELECT s.*, c.name as class_name
        FROM students s
        LEFT JOIN classes c ON s.class_id = c.id
        WHERE {where_clause}
        ORDER BY s.full_name
    """
    
    results = db.conn.execute(query, params).fetchall()
    
    # Convert to Student objects
    students = []
    for row in results:
        student = Student(
            id=row[0], full_name=row[1], birth_date=row[2], address=row[3],
            email=row[4], admission_date=row[5], class_id=row[6],
            profile_image=row[7], gender=row[8], phone=row[9], year=row[10],
            parent_name=row[11], decision_number=row[12], nha_chu_t_info=row[13],
            health_on_admission=row[14], 
            initial_characteristics=row[15] if len(row) > 15 else None
        )
        students.append(student)
    
    return students

def search_veterans_advanced(db, name_query, address, email, gender, health_status):
    """Advanced search for veterans"""
    conditions = []
    params = []
    
    if name_query:
        conditions.append("full_name LIKE ?")
        params.append(f"%{name_query}%")
    
    if address:
        conditions.append("address LIKE ?")
        params.append(f"%{address}%")
        
    if email:
        conditions.append("email LIKE ?")
        params.append(f"%{email}%")
        
    if gender and gender != "Tất cả":
        conditions.append("gender = ?")
        params.append(gender)
        
    if health_status and health_status != "Tất cả":
        conditions.append("initial_characteristics LIKE ?")
        params.append(f"%{health_status}%")
    
    where_clause = " AND ".join(conditions) if conditions else "1=1"
    
    query = f"""
        SELECT * FROM veterans
        WHERE {where_clause}
        ORDER BY full_name
    """
    
    results = db.conn.execute(query, params).fetchall()
    
    # Convert to Veteran objects
    veterans = []
    for row in results:
        veteran = Veteran(
            id=row[0], full_name=row[1], birth_date=row[2], 
            gender=row[3] if len(row) > 3 else None,
            address=row[4] if len(row) > 4 else None,
            email=row[5] if len(row) > 5 else None,
            health_condition=row[6] if len(row) > 6 else None,
            admission_date=row[7] if len(row) > 7 else None,
            psychological_status=row[8] if len(row) > 8 else None,
            profile_image=row[9] if len(row) > 9 else None,
            initial_characteristics=row[10] if len(row) > 10 else None
        )
        veterans.append(veteran)
    
    return veterans

def render_document_management_section(db):
    """Render document management section with instructions"""
    st.subheader("📁 Quản lý tài liệu học sinh")
    
    # Instructions section
    with st.expander("📚 Hướng dẫn sử dụng", expanded=False):
        st.markdown("""
        ### 🔧 Cách sử dụng hệ thống quản lý tài liệu:
        
        **1. Tải lên tài liệu:**
        - Chọn học sinh từ danh sách
        - Chọn file tài liệu (Word, PDF, hình ảnh)
        - Nhập mô tả cho tài liệu
        - Nhấn "Tải lên" để lưu
        
        **2. Xem tài liệu:**
        - Chọn học sinh để xem danh sách tài liệu
        - Nhấn vào tên file để tải xuống
        - Xem thông tin chi tiết: ngày tải, người tải, mô tả
        
        **3. Quản lý tài liệu:**
        - Xóa tài liệu không cần thiết
        - Cập nhật mô tả tài liệu
        - Tìm kiếm tài liệu theo tên hoặc mô tả
        
        **4. Lưu ý quan trọng:**
        - Chỉ tải lên các file liên quan đến học sinh
        - Đặt tên file rõ ràng, dễ hiểu
        - Định kỳ kiểm tra và dọn dẹp tài liệu cũ
        - Kích thước file tối đa: 10MB
        """)
    
    # Check permissions
    user_role = st.session_state.user.role
    if user_role not in ['admin', 'teacher', 'doctor', 'nurse']:
        st.error("Bạn không có quyền truy cập chức năng quản lý tài liệu")
        return
    
    # Student selection
    students = db.get_students()
    if not students:
        st.warning("Không có học sinh nào trong hệ thống")
        return
    
    student_options = [(s.id, f"{s.full_name} (ID: {s.id})") for s in students]
    selected_student_idx = st.selectbox(
        "Chọn học sinh:",
        options=range(len(student_options)),
        format_func=lambda i: student_options[i][1],
        key="doc_student_select"
    )
    
    if selected_student_idx is not None:
        selected_student_id = student_options[selected_student_idx][0]
        selected_student_name = student_options[selected_student_idx][1]
        
        # Create tabs for upload and view
        doc_tabs = st.tabs(["📤 Tải lên tài liệu", "📂 Xem tài liệu"])
        
        with doc_tabs[0]:
            # Upload document form
            st.subheader(f"📤 Tải tài liệu cho {selected_student_name}")
            
            with st.form("upload_document_form"):
                uploaded_file = st.file_uploader(
                    "Chọn tài liệu",
                    type=['pdf', 'doc', 'docx', 'jpg', 'jpeg', 'png', 'txt'],
                    help="Chỉ chấp nhận file PDF, Word, hình ảnh và text"
                )
                
                description = st.text_area(
                    "Mô tả tài liệu",
                    placeholder="Nhập mô tả chi tiết về tài liệu này...",
                    height=100
                )
                
                submit_upload = st.form_submit_button("📤 Tải lên tài liệu", type="primary")
                
                if submit_upload and uploaded_file is not None:
                    if len(description.strip()) < 5:
                        st.error("Vui lòng nhập mô tả chi tiết ít nhất 5 ký tự")
                    else:
                        # Check file size (10MB limit)
                        if uploaded_file.size > 10 * 1024 * 1024:
                            st.error("File quá lớn! Kích thước tối đa là 10MB")
                        else:
                            success = upload_document(db, selected_student_id, uploaded_file, description)
                            if success:
                                st.success(f"✅ Đã tải lên tài liệu '{uploaded_file.name}' thành công!")
                                st.rerun()
                            else:
                                st.error("❌ Không thể tải lên tài liệu. Vui lòng thử lại!")
        
        with doc_tabs[1]:
            # View documents
            st.subheader(f"📂 Tài liệu của {selected_student_name}")
            
            documents = get_student_documents(db, selected_student_id)
            
            if documents:
                st.info(f"Tìm thấy {len(documents)} tài liệu")
                
                # Search documents
                search_term = st.text_input("🔍 Tìm kiếm tài liệu:", placeholder="Nhập tên file hoặc mô tả...")
                
                if search_term:
                    documents = [doc for doc in documents if 
                               search_term.lower() in doc['file_name'].lower() or 
                               search_term.lower() in doc['description'].lower()]
                
                for doc in documents:
                    with st.container():
                        col1, col2, col3 = st.columns([3, 1, 1])
                        
                        with col1:
                            st.markdown(f"**📄 {doc['file_name']}**")
                            st.caption(f"📅 {doc['upload_date']} | 👤 {doc['uploader_name']}")
                            if doc['description']:
                                st.markdown(f"💬 {doc['description']}")
                        
                        with col2:
                            if st.button("📥 Tải xuống", key=f"download_{doc['id']}"):
                                download_document(db, doc)
                        
                        with col3:
                            if user_role in ['admin', 'teacher']:
                                if st.button("🗑️ Xóa", key=f"delete_{doc['id']}"):
                                    if delete_document(db, doc['id']):
                                        st.success("Đã xóa tài liệu!")
                                        st.rerun()
                                    else:
                                        st.error("Không thể xóa tài liệu!")
                        
                        st.divider()
            else:
                st.info("📭 Chưa có tài liệu nào cho học sinh này")

def upload_document(db, student_id, uploaded_file, description):
    """Upload a document for a student"""
    try:
        from datetime import datetime
        
        file_data = uploaded_file.read()
        file_type = uploaded_file.type
        uploader_id = st.session_state.user.id
        
        cursor = db.conn.cursor()
        cursor.execute("""
            INSERT INTO document_files (
                student_id, file_name, file_type, file_data, 
                upload_date, uploaded_by, description
            ) VALUES (?, ?, ?, ?, ?, ?, ?)
        """, (
            student_id, uploaded_file.name, file_type, file_data,
            datetime.now(), uploader_id, description
        ))
        
        db.conn.commit()
        return True
    except Exception as e:
        st.error(f"Lỗi khi tải lên: {str(e)}")
        return False

def get_student_documents(db, student_id):
    """Get all documents for a student"""
    try:
        cursor = db.conn.cursor()
        cursor.execute("""
            SELECT df.id, df.file_name, df.file_type, df.upload_date, 
                   df.description, u.full_name as uploader_name
            FROM document_files df
            LEFT JOIN users u ON df.uploaded_by = u.id
            WHERE df.student_id = ?
            ORDER BY df.upload_date DESC
        """, (student_id,))
        
        documents = []
        for row in cursor.fetchall():
            documents.append({
                'id': row[0],
                'file_name': row[1],
                'file_type': row[2],
                'upload_date': row[3],
                'description': row[4] or "Không có mô tả",
                'uploader_name': row[5] or "Không xác định"
            })
        
        return documents
    except Exception as e:
        st.error(f"Lỗi khi tải danh sách tài liệu: {str(e)}")
        return []

def download_document(db, doc):
    """Download a document"""
    try:
        # Get file data from database
        cursor = db.conn.cursor()
        cursor.execute("SELECT file_data FROM document_files WHERE id = ?", (doc['id'],))
        result = cursor.fetchone()
        
        if result:
            file_data = result[0]
            st.download_button(
                label=f"📥 Tải {doc['file_name']}",
                data=file_data,
                file_name=doc['file_name'],
                mime=doc['file_type'],
                key=f"download_btn_{doc['id']}"
            )
        else:
            st.error("Không tìm thấy file!")
    except Exception as e:
        st.error(f"Lỗi khi tải xuống: {str(e)}")

def delete_document(db, doc_id):
    """Delete a document"""
    try:
        cursor = db.conn.cursor()
        cursor.execute("DELETE FROM document_files WHERE id = ?", (doc_id,))
        db.conn.commit()
        return True
    except Exception as e:
        st.error(f"Lỗi khi xóa tài liệu: {str(e)}")
        return False
    
def render_export_section(db):
    """Render data export section"""
    st.subheader("📤 Xuất dữ liệu")
    
    st.info("📋 Chọn các loại dữ liệu bạn muốn xuất ra file Excel theo định dạng chuẩn")
    
    export_type = st.multiselect(
        "Chọn loại dữ liệu cần xuất:",
        ["Học sinh", "Cựu chiến binh", "Hồ sơ y tế", "Đánh giá tâm lý"],
        default=["Học sinh"]
    )
    
    # T-house selection for students (simplified)
    t_house_option = None
    if "Học sinh" in export_type:
        t_house_options = ["Nhà T2 (mặc định)", "Nhà T3", "Nhà T4", "Tất cả các nhà"]
        
        t_house_option = st.selectbox(
            "Chọn nhà T để xuất:",
            t_house_options,
            index=0
        )
    
    if st.button("📊 Tạo file Excel", type="primary"):
        if not export_type:
            st.warning("Vui lòng chọn ít nhất một loại dữ liệu để xuất")
            return
            
        try:
            with st.spinner("Đang tạo file Excel..."):
                # Create Excel writer
                output = BytesIO()
                from openpyxl import Workbook
                from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
                from openpyxl.utils.dataframe import dataframe_to_rows
                
                wb = Workbook()
                
                # Export students with proper formatting
                if "Học sinh" in export_type:
                    # Create or get worksheet
                    if wb.worksheets:
                        ws = wb.active
                        ws.title = "DANH SÁCH TRẺ EM"
                    else:
                        ws = wb.create_sheet("DANH SÁCH TRẺ EM")
                    
                    # Build SQL query with house filtering
                    house_condition = ""
                    if t_house_option and t_house_option != "Tất cả các nhà":
                        if t_house_option == "Nhà T2 (mặc định)":
                            house_condition = "WHERE s.nha_chu_t_info LIKE '%T2%'"
                        elif t_house_option == "Nhà T3":
                            house_condition = "WHERE s.nha_chu_t_info LIKE '%T3%'"
                        elif t_house_option == "Nhà T4":
                            house_condition = "WHERE s.nha_chu_t_info LIKE '%T4%'"
                        elif t_house_option == "Nhà T5":
                            house_condition = "WHERE s.nha_chu_t_info LIKE '%T5%'"
                        elif t_house_option == "Nhà T6":
                            house_condition = "WHERE s.nha_chu_t_info LIKE '%T6%'"

                    # Get student data filtered by house
                    students = db.conn.execute(f"""
                        SELECT 
                            ROW_NUMBER() OVER (ORDER BY s.id) as stt,
                            s.full_name, 
                            CASE 
                                WHEN s.birth_date IS NOT NULL THEN DATE(s.birth_date)
                                ELSE ''
                            END as birth_date,
                            COALESCE(s.gender, '') as gender,
                            COALESCE(s.phone, '') as phone,
                            COALESCE(s.address, '') as address,
                            COALESCE(s.email, '') as email,
                            CASE 
                                WHEN s.admission_date IS NOT NULL THEN DATE(s.admission_date)
                                ELSE ''
                            END as admission_date,
                            COALESCE(c.name, 'Chưa phân lớp') as class_name,
                            COALESCE(s.year, '') as year,
                            COALESCE(s.parent_name, '') as parent_name,
                            COALESCE(s.decision_number, '') as decision_number,
                            COALESCE(s.nha_chu_t_info, '') as nha_chu_t_info,
                            COALESCE(s.health_on_admission, '') as health_on_admission,
                            COALESCE(s.initial_characteristics, '') as initial_characteristics
                        FROM students s
                        LEFT JOIN classes c ON s.class_id = c.id
                        {house_condition}
                        ORDER BY s.id
                    """).fetchall()
                    
                    # Header formatting
                    header_font = Font(bold=True, size=14)
                    header_alignment = Alignment(horizontal='center', vertical='center')
                    
                    # Determine T-house filter and title
                    if t_house_option and t_house_option != "Tất cả các nhà":
                        if t_house_option == "Nhà T2 (mặc định)":
                            t_house_filter = "NHÀ T2"
                        elif t_house_option == "Nhà T3":
                            t_house_filter = "NHÀ T3"
                        elif t_house_option == "Nhà T4":
                            t_house_filter = "NHÀ T4"
                        elif t_house_option == "Nhà T5":
                            t_house_filter = "NHÀ T5"
                        elif t_house_option == "Nhà T6":
                            t_house_filter = "NHÀ T6"
                        else:
                            t_house_filter = "NHÀ T2"
                    else:
                        t_house_filter = "TẤT CẢ CÁC NHÀ"
                    
                    # Add title with T-house information
                    ws.merge_cells('A1:T1')  # Expanded to cover all columns
                    ws['A1'] = f"DANH SÁCH TRẺ EM {t_house_filter}"
                    ws['A1'].font = Font(bold=True, size=16)
                    ws['A1'].alignment = header_alignment
                    
                    # Set responsible person based on T-house
                    responsible_names = {
                        "NHÀ T2": "Đỗ Khánh Linh",
                        "NHÀ T3": "Chưa cập nhật",
                        "NHÀ T4": "Chưa cập nhật",
                        "NHÀ T5": "Chưa cập nhật",
                        "NHÀ T6": "Chưa cập nhật",
                        "TẤT CẢ CÁC NHÀ": "Nhiều người phụ trách"
                    }
                    responsible_name = responsible_names.get(t_house_filter, "Chưa cập nhật")
                    
                    # Add responsible person info
                    ws.append([f"Mẹ phụ trách : {responsible_name}"])
                    
                    # Add empty row
                    ws.append([])
                    
                    # Add comprehensive headers for ALL fields
                    headers = ['TT', 'Họ và tên', 'Ngày sinh', 'Giới tính', 'Điện thoại', 'Địa chỉ', 'Email', 
                              'Ngày nhập học', 'Lớp', 'Năm', 'Tên bố mẹ/người giám hộ', 'Số quyết định', 'Nhà',
                              'Sức khỏe khi vào làng', 'Đặc điểm sơ bộ']
                    ws.append(headers)
                    
                    # Format header row (now on row 4 due to added responsible person row)
                    for col, header in enumerate(headers, 1):
                        cell = ws.cell(row=4, column=col)
                        cell.font = header_font
                        cell.alignment = header_alignment
                        cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
                    
                    # Add data
                    for student in students:
                        ws.append(list(student))
                    
                    # Set column widths for ALL fields (removed status fields)
                    column_widths = [5, 25, 12, 8, 15, 30, 20, 12, 15, 8, 25, 15, 8, 25, 30]
                    # TT, Name, DOB, Gender, Phone, Address, Email, Admission, Class, Year, Parent, Decision, House, Health_Admission, Characteristics
                    for i, width in enumerate(column_widths, 1):
                        if i <= 26:  # Excel column limit A-Z
                            column_letter = chr(64 + i)  # Convert to A, B, C, etc.
                            ws.column_dimensions[column_letter].width = width
                    
                    st.success(f"✅ Đã xuất {len(students)} học sinh")
                
                # Export veterans
                if "Cựu chiến binh" in export_type:
                    ws_veterans = wb.create_sheet("DANH SÁCH CỰU CHIẾN BINH")
                    
                    # Get ALL veteran data with complete fields
                    veterans = db.conn.execute("""
                        SELECT 
                            ROW_NUMBER() OVER (ORDER BY id) as stt,
                            full_name, 
                            CASE 
                                WHEN birth_date IS NOT NULL THEN DATE(birth_date)
                                ELSE ''
                            END as birth_date,
                            COALESCE(service_period, '') as service_period,
                            COALESCE(health_condition, '') as health_condition,
                            COALESCE(address, '') as address,
                            COALESCE(email, '') as email,
                            COALESCE(contact_info, '') as contact_info,
                            COALESCE(initial_characteristics, '') as initial_characteristics
                        FROM veterans
                        ORDER BY id
                    """).fetchall()
                    
                    # Add title
                    ws_veterans.merge_cells('A1:I1')  # Extended for all columns
                    ws_veterans['A1'] = "DANH SÁCH CỰU CHIẾN BINH"
                    ws_veterans['A1'].font = Font(bold=True, size=16)
                    ws_veterans['A1'].alignment = Alignment(horizontal='center', vertical='center')
                    
                    # Add empty row
                    ws_veterans.append([])
                    
                    # Add comprehensive headers for ALL veteran fields
                    vet_headers = ['TT', 'Họ và tên', 'Ngày sinh', 'Thời gian phục vụ', 
                                  'Sức khỏe khi vào làng', 'Địa chỉ', 'Email', 'Thông tin liên hệ', 'Đặc điểm sơ bộ']
                    ws_veterans.append(vet_headers)
                    
                    # Format header row
                    for col, header in enumerate(vet_headers, 1):
                        cell = ws_veterans.cell(row=3, column=col)
                        cell.font = Font(bold=True, size=14)
                        cell.alignment = Alignment(horizontal='center', vertical='center')
                        cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
                    
                    # Add data
                    for veteran in veterans:
                        ws_veterans.append(list(veteran))
                    
                    # Set column widths for ALL veteran fields
                    vet_column_widths = [5, 25, 15, 20, 25, 30, 20, 20, 35]  # TT, Name, DOB, Service, Health, Address, Email, Contact, Characteristics
                    for i, width in enumerate(vet_column_widths, 1):
                        column_letter = chr(64 + i)  # Convert to A, B, C, etc.
                        ws_veterans.column_dimensions[column_letter].width = width
                    
                    st.success(f"✅ Đã xuất {len(veterans)} cựu chiến binh")
                
                # Export medical records
                if "Hồ sơ y tế" in export_type:
                    ws_medical = wb.create_sheet("HỒ SƠ Y TẾ")
                    
                    medical_records = db.conn.execute("""
                        SELECT 
                            ROW_NUMBER() OVER (ORDER BY mr.date DESC) as stt,
                            CASE 
                                WHEN mr.patient_type = 'student' THEN s.full_name 
                                WHEN mr.patient_type = 'veteran' THEN v.full_name 
                            END as patient_name,
                            mr.patient_type as patient_type,
                            DATE(mr.date) as exam_date,
                            COALESCE(mr.diagnosis, '') as diagnosis,
                            COALESCE(mr.treatment, '') as treatment,
                            COALESCE(mr.medications, '') as medications,
                            CASE 
                                WHEN mr.follow_up_date IS NOT NULL THEN DATE(mr.follow_up_date)
                                ELSE ''
                            END as follow_up_date,
                            COALESCE(u.full_name, '') as doctor_name
                        FROM medical_records mr
                        LEFT JOIN students s ON mr.patient_id = s.id AND mr.patient_type = 'student'
                        LEFT JOIN veterans v ON mr.patient_id = v.id AND mr.patient_type = 'veteran'
                        LEFT JOIN users u ON mr.doctor_id = u.id
                        ORDER BY mr.date DESC
                    """).fetchall()
                    
                    # Add title and headers
                    ws_medical.merge_cells('A1:I1')
                    ws_medical['A1'] = "HỒ SƠ Y TẾ"
                    ws_medical['A1'].font = Font(bold=True, size=16)
                    ws_medical['A1'].alignment = Alignment(horizontal='center', vertical='center')
                    
                    ws_medical.append([])
                    
                    med_headers = ['TT', 'Tên bệnh nhân', 'Loại', 'Ngày khám', 'Chẩn đoán', 
                                  'Điều trị', 'Thuốc', 'Ngày tái khám', 'Bác sĩ']
                    ws_medical.append(med_headers)
                    
                    # Format headers
                    for col, header in enumerate(med_headers, 1):
                        cell = ws_medical.cell(row=3, column=col)
                        cell.font = Font(bold=True, size=14)
                        cell.alignment = Alignment(horizontal='center', vertical='center')
                        cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
                    
                    # Add data
                    for record in medical_records:
                        ws_medical.append(list(record))
                    
                    # Set fixed column widths for medical records
                    med_column_widths = [5, 25, 12, 15, 25, 25, 20, 15, 20]  # TT, Patient, Type, Date, Diagnosis, Treatment, Meds, Follow-up, Doctor
                    for i, width in enumerate(med_column_widths, 1):
                        column_letter = chr(64 + i)  # Convert to A, B, C, etc.
                        ws_medical.column_dimensions[column_letter].width = width
                    
                    st.success(f"✅ Đã xuất {len(medical_records)} hồ sơ y tế")
                
                # Export psychological evaluations
                if "Đánh giá tâm lý" in export_type:
                    ws_psych = wb.create_sheet("ĐÁNH GIÁ TÂM LÝ")
                    
                    psych_evals = db.conn.execute("""
                        SELECT 
                            ROW_NUMBER() OVER (ORDER BY pe.evaluation_date DESC) as stt,
                            s.full_name as student_name,
                            DATE(pe.evaluation_date) as evaluation_date,
                            COALESCE(pe.assessment, '') as assessment,
                            COALESCE(pe.recommendations, '') as recommendations,
                            COALESCE(pe.notes, '') as notes,
                            COALESCE(u.full_name, '') as evaluator_name
                        FROM psychological_evaluations pe
                        LEFT JOIN students s ON pe.student_id = s.id
                        LEFT JOIN users u ON pe.evaluator_id = u.id
                        ORDER BY pe.evaluation_date DESC
                    """).fetchall()
                    
                    # Add title and headers
                    ws_psych.merge_cells('A1:G1')
                    ws_psych['A1'] = "ĐÁNH GIÁ TÂM LÝ"
                    ws_psych['A1'].font = Font(bold=True, size=16)
                    ws_psych['A1'].alignment = Alignment(horizontal='center', vertical='center')
                    
                    ws_psych.append([])
                    
                    psych_headers = ['TT', 'Tên học sinh', 'Ngày đánh giá', 'Đánh giá', 
                                    'Khuyến nghị', 'Ghi chú', 'Người đánh giá']
                    ws_psych.append(psych_headers)
                    
                    # Format headers
                    for col, header in enumerate(psych_headers, 1):
                        cell = ws_psych.cell(row=3, column=col)
                        cell.font = Font(bold=True, size=14)
                        cell.alignment = Alignment(horizontal='center', vertical='center')
                        cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
                    
                    # Add data
                    for eval_record in psych_evals:
                        ws_psych.append(list(eval_record))
                    
                    # Set fixed column widths for psychological evaluations
                    psych_column_widths = [5, 25, 15, 30, 30, 25, 20]  # TT, Student, Date, Assessment, Recommendations, Notes, Evaluator
                    for i, width in enumerate(psych_column_widths, 1):
                        column_letter = chr(64 + i)  # Convert to A, B, C, etc.
                        ws_psych.column_dimensions[column_letter].width = width
                    
                    st.success(f"✅ Đã xuất {len(psych_evals)} đánh giá tâm lý")
                
                # Save workbook to BytesIO
                wb.save(output)
                
                # Create download button
                filename = f"danh_sach_lang_huu_nghi_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
                st.download_button(
                    label="⬇️ Tải xuống file Excel",
                    data=output.getvalue(),
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key="download_excel_export"
                )
                st.success("✅ File Excel đã được tạo thành công theo định dạng chuẩn!")
                
        except Exception as e:
            st.error(f"❌ Lỗi khi tạo file Excel: {str(e)}")
            print(f"Export error: {str(e)}")  # For debugging

if __name__ == "__main__":
    render()